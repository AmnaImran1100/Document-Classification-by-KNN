import os
import docx
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import string

# Download NLTK stopwords if not already downloaded
nltk.download('stopwords')
nltk.download('punkt')

# Function to remove stopwords and punctuation from text
def remove_stopwords_and_punctuation(text):
    stop_words = set(stopwords.words('english'))
    word_tokens = word_tokenize(text)
    filtered_text = [word for word in word_tokens if word.lower() not in stop_words and word not in string.punctuation]
    return ' '.join(filtered_text)

# Process .docx files in a directory and save processed files into a new folder
def process_files(input_folder, output_folder):
    # Create the output folder if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Process each .docx file in the input folder
    for filename in os.listdir(input_folder):
        if filename.endswith('.docx'):
            input_path = os.path.join(input_folder, filename)
            output_path = os.path.join(output_folder, filename)

            # Open the .docx file
            doc = docx.Document(input_path)

            # Extract text from the document
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Remove stopwords and punctuation from the text
            processed_text = remove_stopwords_and_punctuation(text)

            # Create a new document
            new_doc = docx.Document()
            new_doc.add_paragraph(processed_text)

            # Save the new document with processed text
            new_doc.save(output_path)

# Main code
input_folder = 'C:\\Users\\Dell\\Documents\\Github\\Graph_Theory_Project_9296\\Travel'  # Relative path from the current directory
output_folder = 'C:\\Users\\Dell\\Documents\\Github\\Graph_Theory_Project_9296\\ProcessedTravel'  # Relative path from the current directory

# Process files
process_files(input_folder, output_folder)
